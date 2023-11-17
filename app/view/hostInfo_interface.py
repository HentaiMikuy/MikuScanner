# coding:utf-8
from PySide6.QtCore import Slot, Qt
from PySide6.QtGui import QPixmap
from PySide6.QtWidgets import QWidget, QVBoxLayout, QLabel
from qfluentwidgets import ComboBox, PushButton, FluentIcon, InfoBar, InfoBarPosition
from docx import Document

from app.components.table_bar import TabInterface
from app.network.HostInformation import HostInformation

style = """
#title {
    font: 42px 'Segoe UI SemiBold', 'Microsoft YaHei SemiBold';
    background-color: transparent;
    color: black;
}
#low_title {
    font: 15px 'Segoe UI SemiBold', 'Microsoft YaHei SemiBold';
    background-color: transparent;
    color: black;
}
#scanFont {
    font: 25px 'Segoe UI SemiBold', 'Microsoft YaHei SemiBold';
    background-color: transparent;
    color: black;
}
"""


class HostInfoInterface(QWidget):

    def __init__(self, parent=None):
        super().__init__(parent=parent)

        self.setObjectName('HostInfoInterface')
        self.vBoxLayout = QVBoxLayout()
        self.setStyleSheet(style)
        self.view = QWidget(self)
        self.view.resize(850, 650)

        # 标题
        title = QLabel('Access to Information', self.view)
        title.setObjectName("title")
        title.move(40, 35)
        low_title = QLabel(
            'Step 2: Select the IP address of the detected surviving host to obtain more details about the host.',
            self.view
        )
        low_title.setObjectName("low_title")
        low_title.move(40, 89)

        # IP
        ip_label = QLabel('IP', self.view)
        ip_label.setObjectName("scanFont")
        ip_label.move(40, 149)
        self.ip_select = ComboBox(self.view)
        self.ip_select.addItems(
            ['test ip: 192.168.1.1','test ip: 192.168.1.2','test ip: 192.168.1.3']
        )
        self.ip_select.setMinimumSize(210, 47)
        self.ip_select.move(40, 186)

        # start button
        get_button = PushButton('Get it!',self.view, FluentIcon.SEARCH)
        get_button.setMinimumHeight(47)
        get_button.move(300, 186)
        get_button.clicked.connect(self.getHostInfo)

        # Report export
        self.export_button = PushButton('Report export', self.view, FluentIcon.SAVE)
        self.export_button.setMinimumHeight(47)
        self.export_button.move(420, 186)
        self.export_button.setEnabled(False)
        self.export_button.clicked.connect(self.exportToWord)

        # 信息视图
        info_label = QLabel('Information', self.view)
        info_label.setObjectName("scanFont")
        info_label.move(40, 253)
        self.infoTabView = TabInterface(self.view)
        self.infoTabView.move(37, 306)

        # 背景图
        backImage = QLabel(self.view)
        backImage.setPixmap(QPixmap('./app/resource/images/header-host.png'))
        backImage.resize(314, 445)
        backImage.setScaledContents(True)
        backImage.move(566, 225)

        # 功能变量
        self.worker = None

        # 导出到报告所需的临时变量
        self.ip = ""
        self.os = ""
        self.ports = ""

    # 列表IP更新
    @Slot(list)
    def update_IpList(self, new_ipList):
        self.ip_select.clear()
        self.ip_select.addItems(new_ipList)

    def getHostInfo(self):
        InfoBar.info(
            title='Probing',
            content='Please wait patiently for the result.',
            orient=Qt.Horizontal,
            isClosable=True,
            position=InfoBarPosition.TOP_RIGHT,
            duration=4000,
            parent=self
        )
        self.worker = HostInformation(self.ip_select.currentText())
        self.worker.results.connect(self.dataShow)
        self.worker.start()

    def dataShow(self, host, os_info, port_info):
        self.export_button.setEnabled(True)
        self.ip = host
        self.os = os_info
        self.ports = port_info
        self.infoTabView.addTab(host, os_info, port_info)

    def exportToWord(self):
        document = Document()
        document.add_heading('Scan results', 0)
        document.add_heading(f'IP: {self.ip}', level=1)
        document.add_heading('OS:', level=2)
        document.add_paragraph(f'{self.os}').bold = True
        # p_os.bold = True
        document.add_heading('Open Ports:', level=2)
        table = document.add_table(rows=1, cols=3)
        table.style = 'Light List Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Port'
        hdr_cells[1].text = 'Status'
        hdr_cells[2].text = 'Service Name'
        for port, status, service_name in self.ports:
            row_cells = table.add_row().cells
            row_cells[0].text = str(port)
            row_cells[1].text = status
            row_cells[2].text = service_name
        document.add_page_break()
        document.save(f'Scan results for {self.ip}.docx')
        InfoBar.success(
            title='Success',
            content='The scan results were exported successfully.',
            orient=Qt.Horizontal,
            isClosable=True,
            position=InfoBarPosition.TOP_RIGHT,
            duration=4000,
            parent=self
        )
        self.export_button.setEnabled(False)

